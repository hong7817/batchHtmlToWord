#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HTML转Word转换器 - 带完整日志功能
功能：将当前文件夹及子文件夹下的HTML文件转换为Word文档，保留目录结构
作者：DeepSeek助手
版本：2.3.1 (修复打包后输入问题)
"""

import os
import sys
import re
import warnings
import subprocess
import logging
import datetime
import traceback
import time
from pathlib import Path
from collections import deque

# 全局变量
VERSION = "2.3.1"
AUTHOR = "DeepSeek Assistant"

class HtmlToWordConverter:
    """HTML转Word转换器主类"""
    
    def __init__(self):
        # 检测是否打包环境
        self.is_frozen = hasattr(sys, 'frozen')
        self.setup_logging()
        self.total_files = 0
        self.success_files = 0
        self.failed_files = []
        self.output_dir = None
        
    def setup_logging(self):
        """设置日志系统"""
        # 创建日志目录
        log_dir = "logs"
        if not os.path.exists(log_dir):
            os.makedirs(log_dir, exist_ok=True)
        
        # 生成带时间戳的日志文件名
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = os.path.join(log_dir, f"html_to_word_{timestamp}.log")
        
        # 配置日志
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file, encoding='utf-8'),
                logging.StreamHandler(sys.stdout)
            ]
        )
        
        self.logger = logging.getLogger(__name__)
        self.log_file = log_file
        
    def setup_environment(self):
        """设置运行环境"""
        warnings.filterwarnings('ignore')
        os.environ['PYTHONIOENCODING'] = 'utf-8'
        
        if sys.platform.startswith('win'):
            try:
                import _locale
                _locale._getdefaultlocale_backup = _locale._getdefaultlocale
                
                def getpreferredencoding(do_setlocale=True):
                    return 'utf-8'
                
                import locale
                locale.getpreferredencoding = getpreferredencoding
                
                try:
                    import io
                    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='ignore')
                    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='ignore')
                except:
                    pass
                    
            except Exception:
                pass
        
        print("=" * 70)
        print(f"HTML转Word转换器 v{VERSION}")
        print(f"作者: {AUTHOR}")
        print("=" * 70)
        print()
        self.logger.info("=" * 70)
        self.logger.info(f"HTML转Word转换器 v{VERSION} 开始运行")
        self.logger.info("=" * 70)
    
    def check_and_install_dependencies(self):
        """检查并安装依赖"""
        dependencies = [
            ('python-docx', 'docx'),
            ('beautifulsoup4', 'bs4'),
            ('chardet', 'chardet'),
            ('lxml', 'lxml'),
        ]
        
        self.logger.info("正在检查依赖...")
        print("正在检查依赖...")
        
        missing_deps = []
        for pip_name, import_name in dependencies:
            try:
                __import__(import_name)
                self.logger.info(f"✓ {pip_name} 已安装")
                print(f"✓ {pip_name} 已安装")
            except ImportError:
                self.logger.warning(f"✗ {pip_name} 未安装")
                print(f"✗ {pip_name} 未安装")
                missing_deps.append(pip_name)
        
        if missing_deps:
            self.logger.info(f"正在安装缺失的依赖: {', '.join(missing_deps)}")
            print(f"\n正在安装缺失的依赖: {', '.join(missing_deps)}")
            try:
                for dep in missing_deps:
                    self.logger.info(f"正在安装 {dep}...")
                    print(f"正在安装 {dep}...")
                    subprocess.check_call([sys.executable, "-m", "pip", "install", dep, "--quiet"])
                self.logger.info("所有依赖安装完成!")
                print("所有依赖安装完成!")
            except Exception as e:
                self.logger.error(f"安装依赖失败: {str(e)}")
                print(f"安装依赖失败: {str(e)}")
                print("请手动安装:")
                print(f"pip install {' '.join(missing_deps)}")
                return False
        
        self.logger.info("所有依赖检查通过!")
        print("所有依赖检查通过!")
        return True
    
    def detect_encoding(self, file_path):
        """检测文件编码"""
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                
                if result['confidence'] > 0.7:
                    encoding = result['encoding'].lower()
                    if encoding == 'gb2312':
                        return 'gb18030'
                    return encoding
            
            for encoding in ['utf-8', 'gbk', 'gb18030', 'big5']:
                try:
                    raw_data.decode(encoding)
                    return encoding
                except:
                    continue
            
            return 'utf-8'
        except:
            return 'utf-8'
    
    def read_html_file(self, file_path):
        """读取HTML文件内容"""
        try:
            # 先尝试UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    if content.startswith('\ufeff'):
                        content = content[1:]
                    return content
            except UnicodeDecodeError:
                pass
            
            encoding = self.detect_encoding(file_path)
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    if content.startswith('\ufeff'):
                        content = content[1:]
                    return content
            except:
                pass
            
            encodings = ['gbk', 'gb18030', 'big5', 'utf-16', 'latin-1']
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc, errors='ignore') as f:
                        return f.read()
                except:
                    continue
            
            with open(file_path, 'rb') as f:
                content = f.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"读取文件失败 {file_path}: {str(e)}")
            return ""
    
    def find_associated_files_folder(self, html_path):
        """查找关联的_files文件夹"""
        html_dir = os.path.dirname(html_path)
        html_name = os.path.splitext(os.path.basename(html_path))[0]
        
        possible_names = [
            f"{html_name}_files",
            f"{html_name}.files",
            html_name,
            "files",
            "images",
            "image",
            "img"
        ]
        
        for name in possible_names:
            folder_path = os.path.join(html_dir, name)
            if os.path.exists(folder_path) and os.path.isdir(folder_path):
                return folder_path
        
        # 查找包含相似名称的文件夹
        for item in os.listdir(html_dir):
            item_path = os.path.join(html_dir, item)
            if os.path.isdir(item_path):
                if html_name in item.lower() and ('file' in item.lower() or 'image' in item.lower()):
                    return item_path
        
        return None
    
    def convert_html_to_word(self, html_path, docx_path, files_folder=None):
        """转换HTML到Word文档"""
        try:
            from bs4 import BeautifulSoup
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.oxml.ns import qn
            
            self.logger.info(f"开始转换: {html_path}")
            
            # 读取HTML
            html_content = self.read_html_file(html_path)
            if not html_content or not html_content.strip():
                self.logger.warning(f"文件内容为空: {html_path}")
                return False
            
            # 创建Word文档
            doc = Document()
            
            # 设置中文字体
            style = doc.styles['Normal']
            style.font.name = '宋体'
            if hasattr(style._element.rPr, 'add_child'):
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            style.font.size = Pt(10.5)
            
            # 设置标题样式
            for i in range(1, 7):
                try:
                    heading_style = doc.styles[f'Heading {i}']
                    heading_style.font.name = '黑体'
                    if hasattr(heading_style._element.rPr, 'add_child'):
                        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            
            # 解析HTML
            soup = BeautifulSoup(html_content, 'lxml')
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 获取body或整个文档
            body = soup.find('body')
            if not body:
                body = soup
            
            # 处理所有元素
            for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'img', 'ul', 'ol', 'table', 'br']):
                if element.name.startswith('h'):
                    # 标题
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        try:
                            heading = doc.add_heading(text, level=min(level, 6))
                            for run in heading.runs:
                                run.font.name = '黑体'
                                if hasattr(run._element, 'rPr'):
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        except:
                            doc.add_paragraph(f"标题{level}: {text}")
                
                elif element.name == 'p':
                    # 段落
                    text = element.get_text(strip=True)
                    if text:
                        para = doc.add_paragraph()
                        
                        # 处理段落内的内容
                        for child in element.children:
                            if child.name == 'img':
                                # 在段落中插入图片
                                self._insert_image(child, doc, html_path, files_folder, para)
                            elif child.name in ['b', 'strong']:
                                run = para.add_run(child.get_text(strip=True))
                                run.bold = True
                            elif child.name in ['i', 'em']:
                                run = para.add_run(child.get_text(strip=True))
                                run.italic = True
                            elif child.name == 'u':
                                run = para.add_run(child.get_text(strip=True))
                                run.underline = True
                            elif child.name == 'a':
                                run = para.add_run(child.get_text(strip=True))
                                run.font.color.rgb = '0000FF'
                                run.underline = True
                            elif isinstance(child, str):
                                text = str(child).strip()
                                if text:
                                    para.add_run(text)
                        else:
                            # 如果没有子元素，直接添加文本
                            para.add_run(text)
                
                elif element.name == 'div':
                    # div容器
                    text = element.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)
                
                elif element.name == 'img':
                    # 独立图片
                    self._insert_image(element, doc, html_path, files_folder)
                
                elif element.name == 'ul':
                    # 无序列表
                    for li in element.find_all('li', recursive=False):
                        text = li.get_text(strip=True)
                        if text:
                            para = doc.add_paragraph(style='List Bullet')
                            para.add_run(text)
                
                elif element.name == 'ol':
                    # 有序列表
                    for i, li in enumerate(element.find_all('li', recursive=False), 1):
                        text = li.get_text(strip=True)
                        if text:
                            para = doc.add_paragraph(style='List Number')
                            para.add_run(text)
                
                elif element.name == 'table':
                    # 表格
                    rows = element.find_all('tr')
                    if rows:
                        # 计算最大列数
                        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                        if max_cols > 0:
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            for i, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for j, cell in enumerate(cells):
                                    if j < max_cols:
                                        table.cell(i, j).text = cell.get_text(strip=True)
                
                elif element.name == 'br':
                    # 换行
                    if doc.paragraphs:
                        doc.paragraphs[-1].add_run().add_break()
            
            # 如果文档为空，添加一些内容
            if len(doc.paragraphs) == 0:
                body_text = body.get_text(strip=True)
                if body_text:
                    lines = body_text.split('\n')
                    for line in lines[:20]:
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line[:500])
            
            # 保存文档
            output_dir = os.path.dirname(docx_path)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            doc.save(docx_path)
            
            # 验证文件
            if os.path.exists(docx_path) and os.path.getsize(docx_path) > 1024:
                self.logger.info(f"转换成功: {html_path} -> {docx_path}")
                return True
            else:
                self.logger.warning(f"生成的文件可能为空: {docx_path}")
                return False
                
        except Exception as e:
            self.logger.error(f"转换失败 {html_path}: {str(e)}")
            self.logger.error(traceback.format_exc())
            return False
    
    def _insert_image(self, img_element, doc, html_path, files_folder, paragraph=None):
        """插入图片"""
        try:
            from docx.shared import Inches
            
            src = img_element.get('src', '')
            if not src or src.startswith('data:'):
                return
            
            # 清理路径
            clean_src = src.split('?')[0].split('#')[0]
            img_name = os.path.basename(clean_src)
            
            if not img_name:
                return
            
            # 查找图片文件
            img_path = None
            html_dir = os.path.dirname(html_path)
            
            # 1. 在_files文件夹中查找
            if files_folder and os.path.exists(files_folder):
                possible_paths = [
                    os.path.join(files_folder, img_name),
                    os.path.join(files_folder, clean_src)
                ]
                
                # 搜索子目录
                for root, dirs, files in os.walk(files_folder):
                    for file in files:
                        if file == img_name or file.lower() == img_name.lower():
                            possible_paths.append(os.path.join(root, file))
                
                for path in possible_paths:
                    if os.path.exists(path):
                        img_path = path
                        break
            
            # 2. 相对于HTML文件的路径
            if not img_path:
                img_path = os.path.join(html_dir, clean_src)
                if not os.path.exists(img_path):
                    img_path = None
            
            # 添加图片
            if img_path and os.path.exists(img_path):
                try:
                    if paragraph:
                        # 在指定段落中添加
                        run = paragraph.add_run()
                        run.add_picture(img_path, width=Inches(3))
                    else:
                        # 创建新段落
                        para = doc.add_paragraph()
                        para.alignment = 1  # 居中
                        run = para.add_run()
                        run.add_picture(img_path, width=Inches(3))
                except Exception as e:
                    self.logger.warning(f"添加图片失败 {src}: {str(e)}")
                    alt_text = img_element.get('alt', img_name)
                    if paragraph:
                        paragraph.add_run(f"[图片: {alt_text}]")
                    else:
                        doc.add_paragraph(f"[图片: {alt_text}]")
            else:
                alt_text = img_element.get('alt', img_name)
                self.logger.warning(f"图片未找到: {src}")
                if paragraph:
                    paragraph.add_run(f"[图片未找到: {alt_text}]")
                else:
                    doc.add_paragraph(f"[图片未找到: {alt_text}]")
                    
        except Exception as e:
            self.logger.error(f"插入图片失败: {str(e)}")
    
    def process_directory(self):
        """处理目录中的所有HTML文件"""
        current_dir = os.getcwd()
        self.logger.info(f"工作目录: {current_dir}")
        print(f"工作目录: {current_dir}")
        
        # 创建输出目录
        parent_dir = os.path.dirname(current_dir)
        self.output_dir = os.path.join(parent_dir, 'word')
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)
            self.logger.info(f"创建输出目录: {self.output_dir}")
            print(f"创建输出目录: {self.output_dir}")
        else:
            self.logger.info(f"使用现有输出目录: {self.output_dir}")
            print(f"使用现有输出目录: {self.output_dir}")
        
        # 开始扫描
        self.logger.info("开始扫描HTML文件...")
        print("\n开始扫描HTML文件...")
        
        for root, dirs, files in os.walk(current_dir):
            # 过滤掉不需要的目录
            dirs[:] = [d for d in dirs if not d.endswith('_files')]
            dirs[:] = [d for d in dirs if d not in ['word', '__pycache__', '.git', 'logs']]
            
            for file in files:
                if file.lower().endswith(('.html', '.htm')):
                    self.total_files += 1
                    
                    html_path = os.path.join(root, file)
                    
                    # 计算相对路径
                    rel_path = os.path.relpath(root, current_dir)
                    if rel_path == '.':
                        rel_path = ''
                    
                    # 创建输出子目录
                    output_subdir = os.path.join(self.output_dir, rel_path)
                    if not os.path.exists(output_subdir):
                        os.makedirs(output_subdir, exist_ok=True)
                    
                    # 生成输出文件名
                    docx_name = os.path.splitext(file)[0] + '.docx'
                    docx_path = os.path.join(output_subdir, docx_name)
                    
                    # 查找关联的_files文件夹
                    files_folder = self.find_associated_files_folder(html_path)
                    
                    # 转换文件
                    print(f"处理: {file}")
                    if self.convert_html_to_word(html_path, docx_path, files_folder):
                        self.success_files += 1
                        print(f"  ✓ 成功")
                    else:
                        self.failed_files.append(html_path)
                        print(f"  ✗ 失败")
        
        return self.total_files, self.success_files, self.output_dir
    
    def generate_summary_report(self):
        """生成转换摘要报告"""
        summary_file = os.path.join("logs", f"conversion_summary_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
        
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write("=" * 70 + "\n")
            f.write("HTML转Word转换结果摘要\n")
            f.write("=" * 70 + "\n\n")
            
            f.write(f"转换时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"工作目录: {os.getcwd()}\n")
            f.write(f"输出目录: {self.output_dir}\n")
            f.write(f"日志文件: {self.log_file}\n\n")
            
            f.write("=" * 70 + "\n")
            f.write("转换统计\n")
            f.write("=" * 70 + "\n")
            f.write(f"扫描到的HTML文件总数: {self.total_files} 个\n")
            f.write(f"成功转换的文件数: {self.success_files} 个\n")
            f.write(f"转换失败的文件数: {len(self.failed_files)} 个\n")
            
            if self.total_files > 0:
                success_rate = self.success_files / self.total_files * 100
                f.write(f"转换成功率: {success_rate:.1f}%\n")
            f.write("\n")
            
            if self.failed_files:
                f.write("=" * 70 + "\n")
                f.write("转换失败的文件列表\n")
                f.write("=" * 70 + "\n")
                for i, failed_file in enumerate(self.failed_files, 1):
                    f.write(f"{i:3d}. {failed_file}\n")
            else:
                f.write("=" * 70 + "\n")
                f.write("所有文件转换成功!\n")
                f.write("=" * 70 + "\n")
        
        self.logger.info(f"生成转换摘要报告: {summary_file}")
        print(f"生成转换摘要报告: {summary_file}")
        
        return summary_file
    
    def display_summary(self):
        """显示转换摘要"""
        print("\n" + "=" * 70)
        print("转换完成!")
        print("=" * 70)
        print(f"扫描到的HTML文件总数: {self.total_files} 个")
        print(f"成功转换的文件数: {self.success_files} 个")
        print(f"转换失败的文件数: {len(self.failed_files)} 个")
        
        if self.total_files > 0:
            success_rate = self.success_files / self.total_files * 100
            print(f"转换成功率: {success_rate:.1f}%")
        
        print(f"输出目录: {self.output_dir}")
        print(f"日志文件: {self.log_file}")
        print("=" * 70)
        
        if self.failed_files:
            print("\n转换失败的文件:")
            print("-" * 40)
            for i, failed_file in enumerate(self.failed_files, 1):
                print(f"{i:3d}. {failed_file}")
            print("-" * 40)
    
    def safe_exit(self, pause_seconds=5):
        """安全退出程序，适合打包环境"""
        try:
            # 生成完成标记文件
            completion_file = "conversion_completed.txt"
            with open(completion_file, "w", encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("HTML转Word转换完成\n")
                f.write("=" * 60 + "\n")
                f.write(f"完成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"转换文件: {self.success_files}/{self.total_files} 个\n")
                f.write(f"输出目录: {self.output_dir}\n")
                f.write(f"日志文件: {self.log_file}\n")
                f.write("=" * 60 + "\n")
            
            self.logger.info(f"生成完成标记文件: {completion_file}")
            
            # 只在打包环境中显示额外信息
            if self.is_frozen:
                print("\n" + "=" * 60)
                print("程序将在5秒后自动关闭...")
                print("或直接关闭此窗口")
                print(f"转换结果已保存到: {completion_file}")
                print("=" * 60)
                
                # 等待指定秒数
                for i in range(pause_seconds, 0, -1):
                    print(f"\r程序将在 {i} 秒后关闭...", end="", flush=True)
                    time.sleep(1)
                print()
            
        except Exception as e:
            self.logger.error(f"安全退出时出错: {str(e)}")
    
    def run(self):
        """运行转换器"""
        self.setup_environment()
        
        if not self.check_and_install_dependencies():
            self.logger.error("依赖安装失败，程序退出。")
            print("依赖安装失败，程序退出。")
            self.safe_exit(3)
            return
        
        print()
        
        try:
            total, success, output_dir = self.process_directory()
            
            # 生成报告
            summary_file = self.generate_summary_report()
            
            # 显示摘要
            self.display_summary()
            
            if total == 0:
                self.logger.warning("未找到任何HTML文件 (.html 或 .htm)")
                print("\n提示: 未找到任何HTML文件 (.html 或 .htm)")
                print("请将脚本放在包含HTML文件的文件夹中运行")
            
        except KeyboardInterrupt:
            self.logger.warning("用户中断操作。")
            print("\n\n用户中断操作。")
            with open("operation_interrupted.txt", "w", encoding='utf-8') as f:
                f.write("用户中断操作")
            time.sleep(2)
        except Exception as e:
            self.logger.error(f"程序运行出错: {str(e)}", exc_info=True)
            print(f"\n程序运行出错: {str(e)}")
            traceback.print_exc()
            with open("error_report.txt", "w", encoding='utf-8') as f:
                f.write(f"错误时间: {datetime.datetime.now()}\n")
                f.write(f"错误信息: {str(e)}\n")
                f.write(traceback.format_exc())
            time.sleep(3)
        finally:
            # 安全退出
            self.safe_exit()

def main():
    """主函数"""
    try:
        converter = HtmlToWordConverter()
        converter.run()
    except Exception as e:
        print(f"程序启动错误: {e}")
        print(traceback.format_exc())
        # 创建错误报告
        with open("startup_error.txt", "w", encoding='utf-8') as f:
            f.write(f"启动错误时间: {datetime.datetime.now()}\n")
            f.write(f"错误信息: {str(e)}\n")
            f.write(traceback.format_exc())
        
        # 等待后退出
        time.sleep(5)

if __name__ == "__main__":
    main()